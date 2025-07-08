#!/usr/bin/env python3
"""
Knights of Columbus State Officers Word Document Generator

This script reads from the Oklahoma Knights directory SQLite database
and generates a Word document with properly formatted officer tables.

Required packages:
pip install python-docx sqlite3

Usage:
python knights_word_generator.py
"""

import sqlite3
import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
# from docx.oxml.ns import nsdecls
# from docx.oxml import parse_xml

class KnightsWordGenerator:
    def __init__(self, db_path="ok_knights_directory.db"):
        self.db_path = db_path
        self.doc = Document()
        self.setup_document_styles()

    def setup_document_styles(self):
        """Configure document-wide styles"""

        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_topmatter(self):
        """Add the top matter to the document"""

        self.doc.add_section(WD_SECTION.CONTINUOUS)
        self.doc.add_heading("Oklahoma Knights of Columbus State Directory", 0)

    def add_section_header(self, section_title: str):
        """
        Adds a new section to the document;
        this is the masterclass for all different sections
        """
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        heading = self.doc.add_heading(section_title, 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[-1].font.size = Pt(20)
    
    # def set_cell_color(self, cell, color_hex):
    #     """Set background color for a table cell"""
    #     shading_elm = parse_xml(
    #         f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    #     )
    #     cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_officer_table(self, officer_data):
        """
        Create a table for a single officer
        """
        
        # Create the table
        table = self.doc.add_table(rows=4, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        #table.style = 'Table Grid'

        # Set the table width
        # ALERT!!!
        # Ok, so according to stackoverflow, this section of code
        # seems to be ignored by Word, but works in other open office
        # implementations such as google docs. In Word, the column width
        # property is ignored for cell widths.

        # I might try to update this later, but for now, do final
        # rendering of the document in google docs so it looks right.
        table.autofit = False
        DEFAULT_CELL_WIDTH = 1.625
        DEFAULT_CELL_HEIGHT = 0.125
        AFTER_PARAGRAPH_SPACE = Pt(1)

        table.columns[0].width = Inches(DEFAULT_CELL_WIDTH + 0.5)
        table.columns[1].width = Inches(DEFAULT_CELL_WIDTH + 0.5)
        table.columns[2].width = Inches(DEFAULT_CELL_WIDTH - 0.75)
        table.columns[3].width = Inches(DEFAULT_CELL_WIDTH - 0.25)
        # table.rows[0].height = Inches(DEFAULT_CELL_HEIGHT)
        # table.rows[1].height = Inches(DEFAULT_CELL_HEIGHT)
        # table.rows[2].height = Inches(DEFAULT_CELL_HEIGHT)
        # table.rows[3].height = Inches(DEFAULT_CELL_HEIGHT)
        ######
        
        # First Row: Officer Title
        row1 = table.rows[0]

        cell = row1.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        role_run = paragraph.add_run(officer_data['role'])
        role_run.font.bold = True
        role_run.font.size = Pt(12)

        # Second Row: Name, Wife, Council, Primary #
        row2 = table.rows[1]

        cell = row2.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['full_name'])

        cell = row2.cells[1]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['wife'])

        cell = row2.cells[2]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['council'])

        cell = row2.cells[3]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['phone'])

        # Third Row: Address
        row3 = table.rows[2]

        cell = row3.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['address'])
        
        # 4th Row: City/State/Zip, email, ??
        row4 = table.rows[3]

        cell = row4.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['city_state_zip'])

        cell = row4.cells[1]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = AFTER_PARAGRAPH_SPACE
        paragraph.add_run(officer_data['email'])

        # Add spacing after the table
        self.doc.add_paragraph()
    
    def _get_state_officers_data(self):
        """
        Query the database to get state officers information
        
        Returns:
            list: List of dictionaries containing officer data
        """
        if not os.path.exists(self.db_path):
            print(f"Database file not found: {self.db_path}")
            return self.get_sample_data()
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Query to get state officers with full information
            query = """
            SELECT * from "StateOfficerView"
            """
            
            cursor.execute(query)
            rows = cursor.fetchall()
            
            officers = []
            for row in rows:
                officer = {
                    'full_name': row[0] or '[Name Not Available]',
                    'wife': row[1] or '',
                    'address': row[2] or '[Address Not Available]',
                    'city_state_zip': row[3] or '[City Not Available]',
                    'phone': self._format_phone(row[4]),
                    'email': row[5] or '[Email Not Available]',
                    'council': str(row[6]) or '[Council Not Available]',
                    'role': row[8] or '[Role Not Available]'
                }
                officers.append(officer)
            
            conn.close()
            return officers
            
        except sqlite3.Error as e:
            print(f"Database error: {e}")
            return self.get_sample_data()
    
    def create_dd_table(self, district_deputy):
        """
        """

        # Create the table
        table = self.doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row = table.rows[0]

        # First Column: District Number
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Inches(1)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_run = paragraph.add_run(district_deputy['number'])
        role_run.font.bold = True

        # Second Column: Name/Address/City-State-Zip
        cell = row.cells[1]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(district_deputy['district_deputy'])
        cell.add_paragraph(district_deputy['address'])
        cell.add_paragraph(district_deputy['city_state_zip'])

        cell = row.cells[2]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #paragraph.add_run(district_deputy['wife'])
        cell.add_paragraph('')
        cell.add_paragraph(district_deputy['email'])

        cell = row.cells[3]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(district_deputy['phone'])

        for council in district_deputy['councils']:
            cell.add_paragraph(council)

        self.doc.add_paragraph()
    
    def _get_dd_data(self):
        """
        Query the database to get state officers information
        
        Returns:
            list: Returns a list of dictionaries of officer data
        """
        if not os.path.exists(self.db_path):
            print(f"Database file not found: {self.db_path}")
            return self.get_sample_data()
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Query to get state officers with full information
            query = """
            SELECT * from "DistrictsView"
            """
            
            cursor.execute(query)
            rows = cursor.fetchall()
            
            dds = []
            for row in rows:
                address = row[2].split('|') if row[2] is not None else ['null','null','null','null']
                councils = row[6].split('|')

                dd = {
                    'number': str(row[0]) or '[ERROR]',
                    'district_deputy': row[1] or '[VACANT]',
                    'address': address[0] or '',
                    'city_state_zip': address[1] + ', ' + address[2] + ' ' + address[3],
                    'phone': self._format_phone(row[3]) or '',
                    'email': row[4] or '',
                    'home_council': str(row[5]) or '',
                    'councils': councils,
                    #'wife': row[7]
                }
                dds.append(dd)
            
            conn.close()
            return dds
            
        except sqlite3.Error as e:
            print(f"Database error: {e}")
    
    def _format_phone(self, phone):
        """Format phone number for display"""
        if not phone:
            return '[Phone Not Available]'
        
        # Remove non-digits
        digits = ''.join(filter(str.isdigit, str(phone)))
        
        # Format as (XXX) XXX-XXXX if 10 digits
        if len(digits) == 10:
            return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
        elif len(digits) == 11 and digits[0] == '1':
            return f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
        else:
            return str(phone)
    
    def get_sample_data(self):
        """Return sample data if database is not available"""
        return [
            {
                'role': 'STATE SECRETARY',
                'full_name': '[Name from Database]',
                'wife': '[from database]',
                'council': '[Council #]',
                'phone': '[Phone Number]',
                'address': '[Street Address]',
                'city_state_zip': '[City, OK Zip]',
                'email': '[email@domain.com]'
            }
        ]
    
    def generate_officer_section(self):
        """
            add data here
        """
        # Report on section enter.
        print("Generating Officer Section...")

        # Get officer data
        officers = self._get_state_officers_data()

        # Panic if no data found; return
        if not officers:
            print('OFFICER DATA NOT FOUND...SKIPPING SECTION...')
            return

        # If data found, create the section
        self.add_section_header("Oklahoma State Council Officers")

        ## Create table for each officer
        for _, officer in enumerate(officers):
            print(f"Adding {officer['role']}: {officer['full_name']}")
            self.create_officer_table(officer)

        ## Insert a Page Break to end the section
        self.doc.add_page_break()

    def generate_dd_section(self):

        print("Generating District Deputy Section")
        dds = self._get_dd_data()
        if not dds:
            print('DD DATA NOT FOUND...SKIPPING SECTION...')
            return
        
        self.add_section_header("District Deputies")
        self.doc.sections[-1].orientation = WD_ORIENTATION.LANDSCAPE

        ## Create table for each DD
        for _, dd in enumerate(dds):
            self.create_dd_table(dd)
    
    def generate_document(self, output_filename="OK_Knights_StateOfficers.docx"):
        """
        Generate the complete Word document
        
        Args:
            output_filename (str): Name of the output file
        """

        # Generate the Document Header matter
        print("Generating Knights of Columbus State Officers Directory...")
        
        # Add header
        self.add_topmatter()

        # SECTION 1 - STATE OFFICERS
        self.generate_officer_section()

        # SECTION 2 - DISTRICT DEPUTIES
        self.generate_dd_section()

        # SECTION 3 - PROGRAM DIRECTORS / CHAIRMEN
        self.add_section_header("Program Directors and Chairmen")

        # SECTION X - PAST STATE DEPUTIES
        self.add_section_header("Past State Deputies")

        # SECTION Y - WIDOWS OF PAST STATE DEPUTIES
        self.add_section_header("Widows of Past State Deputies")

        # SECTION Z - STATE OFFICER WIVES
        self.add_section_header("Wives of the State Council")
        
        # Save document
        self.doc.save(output_filename)
        print(f"Document saved as: {output_filename}")
        return output_filename

def main():
    """Main function to run the generator"""
    # Initialize generator
    generator = KnightsWordGenerator()
    
    # Generate the document
    output_file = generator.generate_document()
    
    print(f"\nSuccess! State Officers directory created: {output_file}")
    print("\nTo use with your database:")
    print("1. Place your 'ok_knights_directory.db' file in the same folder as this script")
    print("2. Run the script again to populate with real data")
    print("3. The script will automatically format phone numbers and addresses")

if __name__ == "__main__":
    main()